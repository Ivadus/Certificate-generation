import os
import pandas as pd
from docx import Document
from docx2pdf import convert
import tkinter as tk
from tkinter import filedialog, messagebox
import logging

logging.basicConfig(filename='certificates.log', level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s', encoding='utf-8')

# Здесь мы задаем названия столбцов из Excel
COLUMN_FAMILIYA = 'Фамилия ученика'
COLUMN_IMYA = 'Имя ученика'
COLUMN_KLASS = 'Класс'
COLUMN_KL_RUK = 'Классный руководитель'


def replace_text_in_paragraph(paragraph, replacements):
    full_text = paragraph.text
    logging.debug(f"Полный текст параграфа: '{full_text}'")

    new_text = full_text
    for key, value in replacements.items():
        if key in new_text:
            logging.debug(f"Найден заполнитель '{key}' в тексте: '{new_text}'")
            new_text = new_text.replace(key, str(value))
            logging.debug(f"Текст после замены '{key}' на '{value}': '{new_text}'")
        else:
            logging.debug(f"Заполнитель '{key}' не найден в тексте: '{new_text}'")

    if new_text != full_text:
        if paragraph.runs:
            original_run = paragraph.runs[0]
            font = original_run.font
            paragraph.clear()
            new_run = paragraph.add_run(new_text)
            new_run.font.name = font.name
            new_run.font.size = font.size
            new_run.font.bold = font.bold
            new_run.font.italic = font.italic
            new_run.font.underline = font.underline
            if font.color and font.color.rgb:
                new_run.font.color.rgb = font.color.rgb
        else:
            paragraph.add_run(new_text)


def generate_certificates(excel_path, template_path, output_dir, to_pdf=False):
    try:
        logging.info("Начало генерации грамот.")
        data = pd.read_excel(excel_path)
        logging.info(f"Загружено {len(data)} строк из Excel.")

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        for index, row in data.iterrows():
            logging.info(f"Обрабатываем строку {index}: {row.to_dict()}")
            doc = Document(template_path)
            replacements = { #А здесь в фигурных скобках - значения в Word
                '{{Фамилия}}': str(row[COLUMN_FAMILIYA]),
                '{{Имя}}': str(row[COLUMN_IMYA]),
                '{{Класс}}': str(row[COLUMN_KLASS]),
                '{{Кл.рук}}': str(row[COLUMN_KL_RUK])
            }

            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)

            output_file_docx = os.path.join(output_dir,
                                            f"Грамота_{row[COLUMN_FAMILIYA]}_{row[COLUMN_IMYA]}_{index}.docx")
            doc.save(output_file_docx)
            logging.info(f"Создана грамота: {output_file_docx}")

            if to_pdf:
                output_file_pdf = os.path.join(output_dir,
                                               f"Грамота_{row[COLUMN_FAMILIYA]}_{row[COLUMN_IMYA]}_{index}.pdf")
                convert(output_file_docx, output_file_pdf)
                logging.info(f"Создана PDF-версия: {output_file_pdf}")

        logging.info("Генерация грамот завершена.")
        messagebox.showinfo("Успех", "Грамоты успешно сгенерированы.")
    except Exception as e:
        logging.error(f"Ошибка при генерации грамот: {e}")
        messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")


def run_gui():
    root = tk.Tk()
    root.title("Генератор грамот")

    def select_excel():
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        excel_entry.delete(0, tk.END)
        excel_entry.insert(0, path)

    def select_template():
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        template_entry.delete(0, tk.END)
        template_entry.insert(0, path)

    def select_output():
        path = filedialog.askdirectory()
        output_entry.delete(0, tk.END)
        output_entry.insert(0, path)

    def start_generation():
        excel = excel_entry.get()
        template = template_entry.get()
        output = output_entry.get()
        to_pdf = pdf_var.get()
        if not excel or not template or not output:
            messagebox.showwarning("Предупреждение", "Пожалуйста, заполните все поля.")
            return
        generate_certificates(excel, template, output, to_pdf)

    tk.Label(root, text="Excel с данными:").grid(row=0, column=0, padx=5, pady=5)
    excel_entry = tk.Entry(root, width=50)
    excel_entry.grid(row=0, column=1, padx=5, pady=5)
    tk.Button(root, text="Выбрать", command=select_excel).grid(row=0, column=2, padx=5, pady=5)

    tk.Label(root, text="Шаблон Word:").grid(row=1, column=0, padx=5, pady=5)
    template_entry = tk.Entry(root, width=50)
    template_entry.grid(row=1, column=1, padx=5, pady=5)
    tk.Button(root, text="Выбрать", command=select_template).grid(row=1, column=2, padx=5, pady=5)

    tk.Label(root, text="Папка для сохранения:").grid(row=2, column=0, padx=5, pady=5)
    output_entry = tk.Entry(root, width=50)
    output_entry.grid(row=2, column=1, padx=5, pady=5)
    tk.Button(root, text="Выбрать", command=select_output).grid(row=2, column=2, padx=5, pady=5)

    pdf_var = tk.BooleanVar()
    tk.Checkbutton(root, text="Конвертировать в PDF", variable=pdf_var).grid(row=3, column=1, pady=5)

    tk.Button(root, text="Запустить", command=start_generation).grid(row=4, column=1, pady=10)

    root.mainloop()


if __name__ == "__main__":
    run_gui()